"""
ResumeParser: Extracts structured candidate data from PDF, DOCX, and TXT files.

Extraction strategy
-------------------
- Email / phone / URL  → strict regex (near-100% recall on well-formed text)
- Candidate name       → spaCy PERSON entity, with first-line fallback
- Skills               → keyword matching against ALL_SKILLS (case-insensitive, word-boundary)
- Education            → section-header detection + keyword matching
- Experience           → section-header detection + NER ORG/DATE entities
- Certifications       → keyword matching against CERTIFICATION_KEYWORDS
- Languages            → curated language list matching
- Tools                → subset of ALL_SKILLS tagged as tools
- Experience years     → regex on common "X years of experience" patterns + date-range calculation
- Completeness score   → weighted presence of key fields (0–100)
- File hash            → SHA-256 for duplicate detection

PDF parsing order: PyMuPDF (fast) → pdfplumber (better for complex layouts)
"""

import hashlib
import logging
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import spacy

from backend.config import settings
from data.skills_db import (
    ALL_SKILLS,
    CERTIFICATION_KEYWORDS,
    EDUCATION_KEYWORDS,
    EXPERIENCE_PATTERNS,
    PROGRAMMING_LANGUAGES,
    TOOLS_GENERAL,
)

logger = logging.getLogger(__name__)

# ── Regex constants ────────────────────────────────────────────────────────────
_EMAIL_RE = re.compile(r"\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b")
_PHONE_RE = re.compile(
    r"(\+?\d{1,3}[\s\-.]?)?\(?\d{3}\)?[\s\-.]?\d{3}[\s\-.]?\d{4}"
)
_URL_RE = re.compile(r"https?://\S+|www\.\S+", re.IGNORECASE)

# Section headers used in most resumes (order matters — first match wins)
_SECTION_HEADERS = {
    "education": re.compile(
        r"(?:^|\n)\s*(?:education|academic|qualification|degree|university|college)",
        re.IGNORECASE,
    ),
    "experience": re.compile(
        r"(?:^|\n)\s*(?:experience|work history|employment|professional background|career)",
        re.IGNORECASE,
    ),
    "skills": re.compile(
        r"(?:^|\n)\s*(?:skills|technical skills|core competencies|competencies|expertise|proficiencies)",
        re.IGNORECASE,
    ),
    "projects": re.compile(
        r"(?:^|\n)\s*(?:projects|personal projects|academic projects|side projects)",
        re.IGNORECASE,
    ),
    "certifications": re.compile(
        r"(?:^|\n)\s*(?:certifications?|certificates?|credentials?|licen[cs]es?|awards?)",
        re.IGNORECASE,
    ),
    "languages": re.compile(
        r"(?:^|\n)\s*(?:languages?|spoken languages?|language proficiency)",
        re.IGNORECASE,
    ),
}

_DATE_RANGE_RE = re.compile(
    r"((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*\d{4}|"
    r"\d{4})\s*[-–—to]+\s*((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)"
    r"[a-z]*\.?\s*\d{4}|present|current|now|\d{4})",
    re.IGNORECASE,
)

HUMAN_LANGUAGES = [
    "english", "spanish", "french", "german", "chinese", "mandarin", "hindi",
    "arabic", "portuguese", "russian", "japanese", "korean", "italian",
    "dutch", "swedish", "norwegian", "danish", "polish", "turkish",
    "vietnamese", "thai", "indonesian", "malay", "tagalog", "urdu",
    "bengali", "tamil", "telugu", "kannada", "malayalam", "gujarati",
    "marathi", "punjabi", "sinhala", "nepali",
]


class ResumeParser:
    """Parse a resume file or raw text into a structured dict."""

    def __init__(self) -> None:
        try:
            self._nlp = spacy.load(settings.SPACY_MODEL)
        except OSError:
            logger.warning("spaCy model not found, downloading '%s'…", settings.SPACY_MODEL)
            import subprocess
            subprocess.run(
                ["python", "-m", "spacy", "download", settings.SPACY_MODEL],
                check=True,
            )
            self._nlp = spacy.load(settings.SPACY_MODEL)

        # Pre-compile skill patterns for fast matching
        self._skill_patterns: List[Tuple[re.Pattern, str]] = [
            (re.compile(r"\b" + re.escape(skill) + r"\b", re.IGNORECASE), skill)
            for skill in ALL_SKILLS
        ]
        self._lang_patterns: List[Tuple[re.Pattern, str]] = [
            (re.compile(r"\b" + re.escape(lang) + r"\b", re.IGNORECASE), lang)
            for lang in HUMAN_LANGUAGES
        ]
        logger.info("ResumeParser initialised with %d skill patterns.", len(self._skill_patterns))

    # ── Public interface ────────────────────────────────────────────────────────

    def parse(self, filepath: str) -> Dict[str, Any]:
        """Parse a resume file (PDF / DOCX / TXT) and return structured data."""
        path = Path(filepath)
        if not path.exists():
            raise FileNotFoundError(f"Resume file not found: {filepath}")

        ext = path.suffix.lower()
        if ext == ".pdf":
            text = self._read_pdf(filepath)
        elif ext == ".docx":
            text = self._read_docx(filepath)
        elif ext == ".txt":
            text = self._read_txt(filepath)
        else:
            raise ValueError(f"Unsupported format '{ext}'. Accepted: .pdf, .docx, .txt")

        result = self._extract_all(text)
        result["filepath"] = str(filepath)
        result["filename"] = path.name
        result["file_hash"] = self._hash_file(filepath)
        return result

    def parse_text(self, text: str, filename: str = "unknown.txt") -> Dict[str, Any]:
        """Parse raw resume text (useful for loading Kaggle dataset rows)."""
        result = self._extract_all(text)
        result["filepath"] = ""
        result["filename"] = filename
        result["file_hash"] = hashlib.sha256(text.encode("utf-8", errors="ignore")).hexdigest()
        return result

    # ── File readers ────────────────────────────────────────────────────────────

    def _read_pdf(self, filepath: str) -> str:
        text = ""
        # Primary: PyMuPDF
        try:
            import fitz  # PyMuPDF
            with fitz.open(filepath) as doc:
                for page in doc:
                    text += page.get_text()
            if text.strip():
                return text
        except Exception as exc:
            logger.debug("PyMuPDF failed (%s), falling back to pdfplumber.", exc)

        # Fallback: pdfplumber (better for tables / complex layouts)
        try:
            import pdfplumber
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as exc:
            logger.error("pdfplumber also failed: %s", exc)
            raise RuntimeError(f"Could not extract text from PDF '{filepath}'.") from exc

        return text

    def _read_docx(self, filepath: str) -> str:
        from docx import Document
        doc = Document(filepath)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also read tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)

    def _read_txt(self, filepath: str) -> str:
        with open(filepath, "r", encoding="utf-8", errors="ignore") as fh:
            return fh.read()

    # ── Core extraction ─────────────────────────────────────────────────────────

    def _extract_all(self, text: str) -> Dict[str, Any]:
        if not text or not text.strip():
            return self._empty_result(text)

        # Run spaCy on first 100 000 chars (library limit)
        doc = self._nlp(text[:100_000])
        sections = self._split_into_sections(text)

        skills = self._extract_skills(text)
        certifications = self._extract_certifications(text)
        education = self._extract_education(text, sections, doc)
        experience_text = sections.get("experience", "")
        experience_years = self._estimate_experience_years(text)
        projects = self._extract_projects(text, sections)
        languages = self._extract_languages(text)
        tools = self._extract_tools(text, skills)

        parsed: Dict[str, Any] = {
            "raw_text": text,
            "candidate_name": self._extract_name(doc, text),
            "email": self._extract_email(text),
            "phone": self._extract_phone(text),
            "skills": skills,
            "education": education,
            "experience": experience_text[:2000] if experience_text else self._fallback_experience(text),
            "experience_years": experience_years,
            "projects": projects,
            "certifications": certifications,
            "languages": languages,
            "tools": tools,
        }
        parsed["completeness_score"] = self._compute_completeness(parsed)
        return parsed

    def _empty_result(self, text: str) -> Dict[str, Any]:
        return {
            "raw_text": text,
            "candidate_name": None,
            "email": None,
            "phone": None,
            "skills": [],
            "education": [],
            "experience": None,
            "experience_years": 0.0,
            "projects": [],
            "certifications": [],
            "languages": [],
            "tools": [],
            "completeness_score": 0.0,
        }

    # ── Field extractors ────────────────────────────────────────────────────────

    def _extract_name(self, doc: Any, text: str) -> Optional[str]:
        # spaCy PERSON entities — prefer names near the top of the document
        for ent in doc.ents:
            if ent.label_ == "PERSON" and ent.start_char < 500:
                name = ent.text.strip()
                if 2 <= len(name.split()) <= 5:
                    return name

        # Fallback: first non-empty line that looks like a name
        for line in text.split("\n")[:10]:
            line = line.strip()
            if (
                line
                and 2 <= len(line.split()) <= 5
                and not _EMAIL_RE.search(line)
                and not _PHONE_RE.search(line)
                and not any(kw in line.lower() for kw in ("resume", "curriculum", "cv", "profile"))
            ):
                return line

        return None

    def _extract_email(self, text: str) -> Optional[str]:
        match = _EMAIL_RE.search(text)
        return match.group().lower() if match else None

    def _extract_phone(self, text: str) -> Optional[str]:
        match = _PHONE_RE.search(text)
        return match.group().strip() if match else None

    def _extract_skills(self, text: str) -> List[str]:
        found: set[str] = set()
        for pattern, skill in self._skill_patterns:
            if pattern.search(text):
                found.add(skill.lower())
        return sorted(found)

    def _extract_tools(self, text: str, skills: List[str]) -> List[str]:
        tool_keywords = {t.lower() for t in TOOLS_GENERAL}
        return [s for s in skills if s.lower() in tool_keywords]

    def _extract_education(self, text: str, sections: Dict[str, str], doc: Any) -> List[str]:
        edu_section = sections.get("education", text)
        found: List[str] = []

        # Degree / institution patterns
        degree_re = re.compile(
            r"\b(?:b\.?tech|m\.?tech|b\.?e|m\.?e|b\.?sc|m\.?sc|ph\.?d|mba|"
            r"bachelor|master|doctorate|associate|b\.?a|m\.?a)\b",
            re.IGNORECASE,
        )
        for match in degree_re.finditer(edu_section):
            start = max(0, match.start() - 20)
            end = min(len(edu_section), match.end() + 80)
            snippet = edu_section[start:end].replace("\n", " ").strip()
            if snippet not in found:
                found.append(snippet)

        # ORG entities likely to be universities/colleges
        for ent in doc.ents:
            if ent.label_ == "ORG":
                label_lower = ent.text.lower()
                if any(kw in label_lower for kw in ("university", "college", "institute", "school", "academy")):
                    entry = ent.text.strip()
                    if entry not in found:
                        found.append(entry)

        return found[:10]  # cap to avoid noise

    def _extract_experience(self, text: str, sections: Dict[str, str]) -> str:
        return sections.get("experience", "")

    def _fallback_experience(self, text: str) -> str:
        # Return a rough excerpt around "experience" keyword
        match = re.search(r"experience", text, re.IGNORECASE)
        if match:
            start = max(0, match.start() - 50)
            return text[start : start + 500].strip()
        return ""

    def _estimate_experience_years(self, text: str) -> float:
        # Direct "X years of experience" patterns
        for pattern in EXPERIENCE_PATTERNS:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                try:
                    return float(match.group(1))
                except (IndexError, ValueError):
                    pass

        # Date-range approach: sum up months from all date ranges found
        total_months = 0
        for match in _DATE_RANGE_RE.finditer(text):
            start_str = match.group(1).strip()
            end_str = match.group(2).strip()
            months = self._months_between(start_str, end_str)
            total_months += months

        if total_months > 0:
            years = round(total_months / 12, 1)
            return min(years, 40.0)  # sanity cap

        return 0.0

    def _months_between(self, start: str, end: str) -> int:
        """Rough month count between two date strings."""
        current_year = 2026
        current_month = 7

        def _parse(s: str) -> Optional[Tuple[int, int]]:
            s = s.strip().lower()
            if s in ("present", "current", "now"):
                return current_year, current_month
            # Year only
            m = re.match(r"^(\d{4})$", s)
            if m:
                return int(m.group(1)), 1
            # Month Year
            months_map = {
                "jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6,
                "jul": 7, "aug": 8, "sep": 9, "oct": 10, "nov": 11, "dec": 12,
            }
            m = re.match(r"([a-z]+)\.?\s*(\d{4})", s)
            if m:
                mon = months_map.get(m.group(1)[:3], 1)
                return int(m.group(2)), mon
            return None

        s = _parse(start)
        e = _parse(end)
        if s and e:
            diff = (e[0] - s[0]) * 12 + (e[1] - s[1])
            return max(0, diff)
        return 0

    def _extract_projects(self, text: str, sections: Dict[str, str]) -> List[str]:
        proj_section = sections.get("projects", "")
        if not proj_section:
            # Try to find project names near "project" keyword
            proj_section = text

        projects: List[str] = []
        # Lines that start with a bullet, dash, number, or capital letter in project section
        for line in proj_section.split("\n"):
            line = line.strip()
            if 10 < len(line) < 200 and re.match(r"^[•\-\*\d]|^[A-Z]", line):
                projects.append(line)
        return projects[:15]

    def _extract_certifications(self, text: str) -> List[str]:
        found: List[str] = []
        for keyword in CERTIFICATION_KEYWORDS:
            pattern = re.compile(
                r"[^\n]*\b" + re.escape(keyword) + r"\b[^\n]*", re.IGNORECASE
            )
            for match in pattern.finditer(text):
                snippet = match.group().strip()
                if 5 < len(snippet) < 200 and snippet not in found:
                    found.append(snippet)
        return found[:10]

    def _extract_languages(self, text: str) -> List[str]:
        found: List[str] = []
        for pattern, lang in self._lang_patterns:
            if pattern.search(text):
                found.append(lang)
        return sorted(set(found))

    # ── Section splitter ────────────────────────────────────────────────────────

    def _split_into_sections(self, text: str) -> Dict[str, str]:
        """
        Split resume text into named sections by detecting common section headers.
        Returns a dict of section_name → section_text.
        """
        # Find positions of all section headers
        hits: List[Tuple[int, str]] = []
        for section_name, pattern in _SECTION_HEADERS.items():
            for match in pattern.finditer(text):
                hits.append((match.start(), section_name))

        if not hits:
            return {}

        hits.sort(key=lambda x: x[0])
        sections: Dict[str, str] = {}
        for i, (pos, name) in enumerate(hits):
            end = hits[i + 1][0] if i + 1 < len(hits) else len(text)
            sections[name] = text[pos:end]

        return sections

    # ── Completeness score ──────────────────────────────────────────────────────

    def _compute_completeness(self, parsed: Dict[str, Any]) -> float:
        weights = {
            "candidate_name": 10,
            "email": 10,
            "phone": 5,
            "skills": 25,
            "education": 15,
            "experience": 15,
            "projects": 10,
            "certifications": 5,
            "languages": 3,
            "tools": 2,
        }
        score = 0.0
        for field, weight in weights.items():
            value = parsed.get(field)
            if value:
                score += weight
        return round(score, 1)

    # ── File hash ───────────────────────────────────────────────────────────────

    def _hash_file(self, filepath: str) -> str:
        h = hashlib.sha256()
        with open(filepath, "rb") as fh:
            for chunk in iter(lambda: fh.read(65_536), b""):
                h.update(chunk)
        return h.hexdigest()
